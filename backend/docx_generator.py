import io

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

from config import SMATA_COLORS

NETWORK_LABELS = {
    "twitter": "X",
    "instagram": "INSTAGRAM",
    "tiktok": "TIK TOK",
}

NETWORK_ORDER = ["twitter", "instagram", "tiktok"]

POST_TEXT_COLOR = RGBColor(0x80, 0x00, 0x00)  # rojo oscuro institucional
LINK_COLOR_HEX = "0563C1"  # azul hyperlink estándar de Word
MAX_POST_CHARS = 500       # truncado conservador para textos muy largos

HYPERLINK_REL = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink"


def _hex(color: str) -> str:
    return color.lstrip("#")


def _add_hyperlink(paragraph, url: str, text: str, color_hex: str = LINK_COLOR_HEX) -> None:
    """Inserta un hyperlink clickeable, azul y subrayado (python-docx no lo soporta nativo)."""
    part = paragraph.part
    r_id = part.relate_to(url, HYPERLINK_REL, is_external=True)

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")

    color = OxmlElement("w:color")
    color.set(qn("w:val"), color_hex)
    r_pr.append(color)

    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)

    run.append(r_pr)

    t = OxmlElement("w:t")
    t.text = text
    t.set(qn("xml:space"), "preserve")
    run.append(t)

    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def _shade_cell(cell, color_hex: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def _truncate(text: str, limit: int = MAX_POST_CHARS) -> str:
    text = (text or "").strip()
    if len(text) <= limit:
        return text
    cut = text[:limit].rsplit(" ", 1)[0]
    return cut + "…"


def _add_title_block(document: Document) -> None:
    """Bloque destacado 'TENDENCIA EN REDES' con sombreado institucional."""
    table = document.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]
    _shade_cell(cell, _hex(SMATA_COLORS["GREEN_DARK"]))
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    para = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after = Pt(6)

    run = para.add_run("TENDENCIA EN REDES")
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    document.add_paragraph()  # respiro debajo del bloque


def _add_network_header(document: Document, label: str) -> None:
    para = document.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    para.paragraph_format.space_before = Pt(12)
    para.paragraph_format.space_after = Pt(6)

    run = para.add_run(label)
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor.from_string(_hex(SMATA_COLORS["GREEN_DARK"]))


def _add_post(document: Document, post: dict) -> None:
    text = _truncate(post.get("text") or "").upper() or "(SIN TEXTO)"

    text_para = document.add_paragraph()
    text_para.paragraph_format.space_after = Pt(2)
    text_run = text_para.add_run(text)
    text_run.font.color.rgb = POST_TEXT_COLOR
    text_run.font.size = Pt(11)

    link_para = document.add_paragraph()
    link_para.paragraph_format.space_after = Pt(10)
    post_url = post.get("post_url") or ""
    if post_url:
        _add_hyperlink(link_para, post_url, post_url)


def generate_docx(posts: list[dict]) -> bytes:
    document = Document()
    _add_title_block(document)

    by_network: dict[str, list[dict]] = {}
    for p in posts:
        net = p.get("network", "")
        by_network.setdefault(net, []).append(p)

    for net in NETWORK_ORDER:
        items = by_network.get(net, [])
        if not items:
            continue
        _add_network_header(document, NETWORK_LABELS[net])
        for post in items:
            _add_post(document, post)

    # Fallback: redes no mapeadas
    known = set(NETWORK_ORDER)
    for net, items in by_network.items():
        if net in known or not items:
            continue
        _add_network_header(document, (net or "OTROS").upper())
        for post in items:
            _add_post(document, post)

    buf = io.BytesIO()
    document.save(buf)
    buf.seek(0)
    return buf.read()
